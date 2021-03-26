from PIL import ImageGrab
from PIL import Image
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime

def cap_screen(x1,y1,x2,y2,path,file_name):
    '''
    :param x1: int
    :param y1: int
    :param x2: int
    :param y2: int
    :param path: str
    :param file_name: str
    :return: None
    '''
    full_path = path  + '/' + file_name
    image = ImageGrab.grab(bbox=(x1,y1,x2,y2))
    image.save(full_path)

def read_text2list(path, file_name, cut='\n\n'):
    '''
    :param path: str
    :param file_name: str
    :param cut: str
    :return: list [str]
    '''
    full_path = path + '/' + file_name
    with open(full_path, 'r') as f:
        text = f.read()
        text_list = text.split(cut)
    text_list = text_list[0:-1]
    return text_list

def append_text(text, path, file_name):
    '''
    :param text: str
    :param path: str
    :param file_name: str
    :return: list [str]
    '''
    full_path = path + '/' + file_name
    with open(full_path, 'a') as f:
        f.write(text + '\n\n')

def compara_transcript_with_mark(transcript_list, mark_list):
    '''
    :param transcript_list: [str]
    :param mark_list: [str]
    :return: [str]
    '''
    i = 0
    text_buffer = ''
    text_list = []
    for sentc in transcript_list:
        text_buffer += sentc + ' '
        if i < len(mark_list):
            if sentc == mark_list[i]:
                text_list.append(text_buffer)
                text_buffer = ''
                i += 1
    # text_list.append(text_buffer)  # dont forget the last append
    return text_list

def init_docs(main_path, img_dir, text_dir):
    imges_path = main_path + '/' + img_dir
    for f in os.listdir(imges_path):
        os.remove(imges_path + '/' + f)
    text_path = main_path + '/' + text_dir
    for f in os.listdir(text_path):
        os.remove(text_path + '/' + f)

def write_notes(text_list, main_path, img_dir, text_dir):
    # current time for naming
    now = datetime.now()
    DT_STRING = now.strftime("%y%m%d%H%M%S")
    load_path = main_path + '/' + img_dir
    saved_path = main_path + '/' + text_dir

    # loop for all file name, then transfer image from png to jpg
    for i, f in enumerate(sorted(os.listdir(load_path))):
        im = Image.open(load_path + '/' + f)
        im = im.convert('RGB')
        t_f = '{}.png'.format(i)
        im.save(load_path + '/' + t_f)

    # write docx including text and pictures
    document = Document()
    p = document.add_paragraph()
    r = p.add_run()
    font = r.font
    font.name = "Calibri"
    font.size = Pt(11)

    # write the txt in doc file
    for i, t in enumerate(text_list):
        r.add_picture(load_path + '/{}.png'.format(i), width=Inches(5.5))
        r.add_break()
        r.add_text(t)
        r.add_break()
        if i != len(text_list) - 1:
            r.add_break()
    document.save(saved_path + '/{}_udemy_notes.docx'.format(DT_STRING))